import io
import re
from typing import Optional
import fitz  # pymupdf
from docx import Document


# ─────────────────────────────────────────────────────────────────────────────
# Language vocabulary
# ─────────────────────────────────────────────────────────────────────────────

# Sorted by length descending so multi-word phrases match before single words.
_LEVEL_KEYWORDS = {
    # Multi-word phrases (checked first due to sort-by-length)
    "native language":                  "Native",
    "native speaker":                   "Native",
    "mother tongue":                    "Native",
    "langue maternelle":                "Native",
    "professional working proficiency": "Professional",
    "full professional proficiency":    "Professional",
    "professional proficiency":         "Professional",
    "advanced proficiency":             "Advanced",
    "upper intermediate":               "Upper Intermediate",
    "limited working proficiency":      "Elementary",
    "elementary proficiency":           "Elementary",
    # Single words
    "native":         "Native",
    "fluent":         "Fluent",
    "professional":   "Professional",
    "advanced":       "Advanced",
    "intermediate":   "Intermediate",
    "elementary":     "Elementary",
    "beginner":       "Beginner",
    "basic":          "Beginner",
    "b2": "B2", "b1": "B1", "a2": "A2", "a1": "A1",
    "bilingue":       "Fluent",
    "biginner":       "Beginner",   # common CV typo
    "maternelle":     "Native",
    "courant":        "Fluent",
    "professionnel":  "Professional",
    "avancé":         "Advanced",
    "avance":         "Advanced",
    "intermédiaire":  "Intermediate",
    "intermediaire":  "Intermediate",
    "débutant":       "Beginner",
    "debutant":       "Beginner",
    "notions":        "Beginner",
}

_LANGUAGE_NAMES = {
    # French → English
    "arabe": "Arabic", "français": "French", "francais": "French",
    "anglais": "English", "espagnol": "Spanish", "allemand": "German",
    "italien": "Italian", "chinois": "Chinese", "japonais": "Japanese",
    "russe": "Russian", "portugais": "Portuguese", "néerlandais": "Dutch",
    "neerlandais": "Dutch", "turc": "Turkish", "coréen": "Korean",
    # English (pass-through)
    "arabic": "Arabic", "french": "French", "english": "English",
    "spanish": "Spanish", "german": "German", "italian": "Italian",
    "chinese": "Chinese", "japanese": "Japanese", "russian": "Russian",
    "portuguese": "Portuguese", "dutch": "Dutch", "turkish": "Turkish",
    "korean": "Korean",
}


# ─────────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith(".doc"):
        raise ValueError("Legacy .doc format not supported. Please use .docx or .pdf")
    else:
        raise ValueError(f"Unsupported file format: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# PDF extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        full_text = []
        for page in doc:
            blocks = page.get_text("blocks")
            blocks_sorted = sorted(blocks, key=lambda b: (round(b[1] / 15), b[0]))
            for block in blocks_sorted:
                text = block[4].strip()
                if text:
                    full_text.append(text)
        doc.close()
        raw = "\n".join(full_text)
        raw = normalize_spaced_text(raw)
        raw = _fix_body_urls(raw)
        return raw
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def normalize_spaced_text(text: str) -> str:
    """
    Collapse CSS letter-spaced text: 'B E N J A M I N' → 'BENJAMIN'.
    Double-spaces mark word boundaries; single-spaces are intra-word.
    Normal prose lines are untouched.
    """
    result_lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            result_lines.append(line)
            continue
        tokens = stripped.split()
        is_spaced = len(tokens) >= 3 and all(len(t) == 1 for t in tokens)
        if is_spaced:
            source = line if "  " in line else stripped
            if "  " in source:
                groups = re.split(r"  +", source.strip())
                normalized = " ".join(g.replace(" ", "") for g in groups)
            else:
                normalized = stripped.replace(" ", "")
            result_lines.append(normalized)
        else:
            result_lines.append(line)
    return "\n".join(result_lines)


def _fix_body_urls(text: str) -> str:
    return re.sub(
        r"(https?://[a-zA-Z0-9.\-/]+)\s*\n\s*([a-zA-Z0-9_\-/][^\s]*)",
        r"\1\2", text,
    )


# ─────────────────────────────────────────────────────────────────────────────
# Contact field extraction
# ─────────────────────────────────────────────────────────────────────────────

_NEXT_BOUNDARY = re.compile(
    r"(?:Phone|Téléphone|Tel|Email|Linkedin|LinkedIn|Github|GitHub|"
    r"Portfolio|Address|Adresse|PROFILE|PROFIL|SKILLS|COMPÉTENCES|"
    r"PROFESSIONAL|EDUCATION|ÉDUCATION|EXTRACURRICULAR|LANGUAGES?|"
    r"LANGUE[S]?|CERTIFICATIONS?)\s*[:\|]",
    re.IGNORECASE,
)


def _chunk_after_label(text: str, label_re: str) -> str:
    m = re.search(label_re, text, re.IGNORECASE)
    if not m:
        return ""
    start = m.end()
    nxt = _NEXT_BOUNDARY.search(text, start)
    end = nxt.start() if nxt else min(start + 200, len(text))
    return re.sub(r"\s+", "", text[start:end])


def extract_contact_fields(text: str) -> dict:
    """
    Extract email, phone, LinkedIn, GitHub from anywhere in the CV text.
    Uses label-anchored search with full-text fallback.
    """
    email_chunk    = _chunk_after_label(text, r"Email\s*:\s*")
    # Phone: search the original text after the label (not collapsed) to preserve spacing
    linkedin_chunk = _chunk_after_label(text, r"Linked[Ii]n\s*:\s*")
    github_chunk   = _chunk_after_label(text, r"[Gg]it[Hh]ub\s*:\s*")

    email_m = re.search(
        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9\-]+(?:\.[a-zA-Z0-9\-]+)*\.[a-zA-Z]{2,6}",
        email_chunk,
    )
    # Phone: always search original text to preserve digit spacing (+216 29 176 273)
    # Use [^\n] to prevent matching across newlines into address numbers
    phone_m    = re.search(r"(\+?\d[\d \-\(\)]{6,}\d)(?!\d)", text)
    linkedin_m = re.search(
        r"(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_\-]+", linkedin_chunk,
    )
    github_m   = re.search(r"https?://github\.com/[a-zA-Z0-9_\-]+", github_chunk)

    if not email_m:
        email_m = re.search(
            r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,6}", text,
        )
    if not linkedin_m:
        linkedin_m = re.search(
            r"(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_\-]+", text,
        )
    if not github_m:
        # Full-text fallback — only match URLs that appear in contact/social
        # context (near "github", "git", or at line start), NOT inside
        # descriptions or body text where other people's URLs might appear
        github_m = re.search(
            r"(?:github\.com/|git\.com/)([a-zA-Z0-9_\-]+)",
            github_chunk or ""
        )
        if not github_m:
            # Last resort: only match if the URL is on its own line or
            # preceded by a contact-like label
            for line in text.split("\n"):
                line = line.strip()
                m = re.search(r"https?://github\.com/([a-zA-Z0-9_\-]+)$", line)
                if m and len(line) < 60:  # short line = likely contact row
                    github_m = m
                    break

    def _https(url: Optional[str]) -> Optional[str]:
        return ("https://" + url) if url and not url.startswith("http") else url

    def _github_url(m) -> Optional[str]:
        if not m:
            return None
        matched = m.group(0)
        # If it's a full URL, use as-is
        if matched.startswith("http"):
            return matched
        # If it's a partial match (github.com/username), prepend https://
        return "https://" + matched if matched.startswith("github") else None

    return {
        "email":    email_m.group(0)            if email_m    else None,
        "phone":    phone_m.group(0).strip()    if phone_m    else None,
        "linkedin": _https(linkedin_m.group(0)) if linkedin_m else None,
        "github":   _github_url(github_m),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Language extraction — deterministic, never hallucinates
# ─────────────────────────────────────────────────────────────────────────────

def extract_languages_from_text(text: str) -> list:
    """
    Deterministically extract languages and proficiency levels.

    Handles all formats:
    - Sidebar: names on one line, levels on next  ('Arabe Francais' / 'NATIVE ADVANCED')
    - Inline with qualifiers: 'Arabic Native language French Professional proficiency'
    - Inline simple: 'Arabic Native  French Professional'
    - No levels: names only → all levels null
    - Parenthetical: 'Arabic (Native)'
    - Dash/pipe separated: 'Arabic - Native'

    Multi-word level phrases ('Native language', 'Professional proficiency') are
    matched before single-word keywords so they are correctly identified.
    Never guesses or infers a level — if none is present, returns null.
    """
    lang_start = re.search(
        r"(?:LANGUAGES?\s*SPOKEN|LANGUES?\s*PARLÉES?|LANGUAGES?|LANGUE[S]?|COMPÉTENCES\s+LINGUISTIQUES?)\s*[:\n]",
        text, re.IGNORECASE,
    )
    if not lang_start:
        # Fallback: look for "Français :" or "French :" pattern directly
        lang_start = re.search(
            r"(?:Français|French|Anglais|English|Arabe|Arabic|Espagnol|Spanish)\s*:",
            text, re.IGNORECASE,
        )
    if not lang_start:
        return []

    next_sec = re.search(
        r"\n(?:PROFILE|PROFIL|SKILLS|COMPÉTENCES|EDUCATION|ÉDUCATION|FORMATION|"
        r"EXPERIENCE|EXPÉRIENCE|CERTIFICATIONS?|AWARDS?|PROJECTS?|HACKATHONS?|"
        r"VOLUNTEER|BÉNÉVOLAT|CONTACT|SUMMARY|RÉSUMÉ)\b",
        text[lang_start.end():], re.IGNORECASE,
    )
    lang_text = text[
        lang_start.end():
        lang_start.end() + (next_sec.start() if next_sec else 600)
    ]

    # Build patterns — longest phrases first so 'native language' beats 'native'
    level_pat = re.compile(
        r"(?:" + "|".join(re.escape(k) for k in
                          sorted(_LEVEL_KEYWORDS, key=len, reverse=True)) + r")",
        re.IGNORECASE,
    )
    lang_pat = re.compile(
        r"\b(" + "|".join(re.escape(k) for k in _LANGUAGE_NAMES) + r")\b",
        re.IGNORECASE,
    )

    def _is_level_only_line(line: str) -> bool:
        """
        Return True if the line contains ONLY level content.
        Handles both single tokens ('ADVANCED') and multi-word phrases
        ('Native language', 'Professional proficiency').
        """
        stripped = line.strip()
        if not stripped:
            return False
        # Remove all level phrases — if nothing non-whitespace remains, it's a level line
        remaining = level_pat.sub("", stripped).strip()
        return len(remaining) == 0

    # ── Strategy A: separate lines (sidebar format) ──────────────────────────
    lines = [l.strip() for l in lang_text.split("\n") if l.strip()]
    name_lines, level_lines = [], []

    for line in lines:
        tokens = re.sub(r"[•\-\|,]", " ", line).split()
        tokens = [t for t in tokens if t]
        if not tokens:
            continue
        if all(lang_pat.fullmatch(t) for t in tokens):
            name_lines.append(tokens)
        elif _is_level_only_line(line):
            # Extract all level phrases from this line in order
            level_values = [m.group(0) for m in level_pat.finditer(line)]
            if level_values:
                level_lines.append(level_values)

    names  = [t for line in name_lines  for t in line]
    levels = [t for line in level_lines for t in line]

    # Strategy A: separate lines (sidebar format)
    # When counts match exactly → pair by index (e.g. Selim's 4-name/4-level sidebar)
    # When counts don't match but both name_lines and level_lines exist →
    #   some names have no level (blank cell in PDF table). Pad with None and pair.
    if name_lines and level_lines:
        if len(levels) != len(names):
            # Pad shorter list with None so index pairing still works
            levels = levels + [None] * (len(names) - len(levels))
        results, seen = [], set()
        for i, name_token in enumerate(names):
            canonical = _LANGUAGE_NAMES.get(name_token.lower())
            if not canonical or canonical in seen:
                continue
            seen.add(canonical)
            level_key = levels[i].lower() if (i < len(levels) and levels[i]) else None
            level = _LEVEL_KEYWORDS.get(level_key) if level_key else None
            results.append({"name": canonical, "level": level})
        if results:
            return results

    # ── Strategy B: inline pairing ───────────────────────────────────────────
    # Find each language name, then scan the next 60 chars for a level phrase.
    # Stop the window at the next language name to avoid cross-pairing.
    results, seen = [], set()
    lang_matches = list(lang_pat.finditer(lang_text))

    for lm in lang_matches:
        canonical = _LANGUAGE_NAMES.get(lm.group(0).lower())
        if not canonical or canonical in seen:
            continue
        seen.add(canonical)

        level = None
        window_text = lang_text[lm.end(): lm.end() + 60]
        # Truncate window at the next language name
        next_lang = lang_pat.search(window_text)
        window_text = window_text[:next_lang.start()] if next_lang else window_text

        lv_match = level_pat.search(window_text)
        if lv_match:
            level = _LEVEL_KEYWORDS.get(lv_match.group(0).lower())

        results.append({"name": canonical, "level": level})

    return results


# ─────────────────────────────────────────────────────────────────────────────
# DOCX extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")